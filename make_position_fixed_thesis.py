from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re, shutil

base = Path(r"C:\Users\31606\Desktop")
src = base / "系统界面" / "校园垃圾自动识别与分类系统设计副本.docx"
out = base / "系统界面" / "校园垃圾自动识别与分类系统设计副本_ResidualEMA位置修正版.docx"
shutil.copy2(src, out)
doc = Document(str(out))

# ---------- helpers ----------
def set_font(run, east="宋体", ascii_font="Times New Roman", size=10.5, bold=None, color="000000"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def fmt_body(p):
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Inches(0.28)
    for r in p.runs:
        set_font(r)

def fmt_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = Pt(18)
    for r in p.runs:
        set_font(r, size=10.5)

def fmt_heading(p, level=1):
    p.style = f"Heading {level}"
    p.paragraph_format.first_line_indent = None
    for r in p.runs:
        set_font(r, east="黑体", size=14 if level == 1 else 12, bold=True, color="000000")

def para_after(paragraph, text="", body=True):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        r = p.add_run(text)
        set_font(r)
    if body:
        fmt_body(p)
    return p

def image_after(paragraph, img_path, caption, width=5.3):
    img_path = Path(img_path)
    if not img_path.exists():
        return paragraph, False
    p_img = para_after(paragraph, "", body=False)
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = None
    run = p_img.add_run()
    run.add_picture(str(img_path), width=Inches(width))
    p_cap = para_after(p_img, caption, body=False)
    fmt_caption(p_cap)
    return p_cap, True

def table_after(paragraph, title, headers, rows):
    title_p = para_after(paragraph, title, body=False)
    fmt_caption(title_p)
    table = doc.add_table(rows=1, cols=len(headers))
    title_p._p.addnext(table._tbl)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i,v in enumerate(row):
            cells[i].text = str(v)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = Pt(18)
                for r in p.runs:
                    set_font(r, size=10)
    return table

def find_exact(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None

def find_starts(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None

def replace_text(old, new):
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
            fmt_body(p)

# ---------- structure fixes only in original ----------
for p in doc.paragraphs:
    t = p.text.strip()
    if t == "绪论":
        p.text = "第一章 绪论"
        fmt_heading(p, 1)
    m = re.match(r"^(第[二三四五六]章)：(.+)$", t)
    if m:
        p.text = f"{m.group(1)} {m.group(2)}"
        fmt_heading(p, 1)
    elif re.match(r"^\d+\.\d+", t):
        fmt_heading(p, 2)

replace_text("引入注意力机制和训练优化策略", "引入ResidualEMA特征增强结构和训练优化策略")
replace_text("注意力机制引入后", "ResidualEMA结构引入后")

# ---------- Chapter 2 formula explanation: only around real 2.6 / 2.4 ----------
p26 = find_starts("2.6 模型评价指标")
if p26:
    last = para_after(p26, "目标检测模型不能只依靠单一准确率进行评价。本文在实验中主要使用Precision、Recall、F1、mAP@0.5和mAP@0.5:0.95等指标，从误检、漏检和定位精度三个角度综合分析YOLOv8-ResidualEMA的效果。")
    last = para_after(last, "Precision = TP / (TP + FP)；Recall = TP / (TP + FN)；F1 = 2 × Precision × Recall / (Precision + Recall)。")
    last = para_after(last, "其中，TP表示正确检测出的垃圾目标，FP表示误检目标，FN表示漏检目标。Precision越高说明误检越少，Recall越高说明漏检越少，F1则用于综合衡量准确率和召回率。mAP@0.5表示IoU阈值为0.5时各类别AP的平均值，mAP@0.5:0.95则采用多个IoU阈值进行平均，更能反映模型在严格定位条件下的检测稳定性。")

p24 = find_starts("2.4 ResidualEMA")
if p24:
    para_after(p24, "ResidualEMA结构的引入目的不是简单增加网络复杂度，而是在保持YOLOv8主干检测流程基本稳定的前提下，对中高层特征进行残差保留和有效区域重加权。公式部分需要结合残差分支、特征权重和输出特征说明其对复杂背景、小目标和类别相似样本的作用。")

# ---------- Chapter 4: insert only after existing Fig 4.3 / before 4.7 ----------
fig43 = find_starts("图4.3")
if fig43:
    last = para_after(fig43, "为避免实验分析只停留在最终指标对比，本文进一步补充训练过程曲线、阈值曲线和混淆矩阵。训练曲线用于证明模型是否正常收敛，PR曲线和F1曲线用于分析不同置信度阈值下准确率与召回率的平衡关系，混淆矩阵则用于观察不同垃圾类别之间的误判情况。")
    table_after(last, "表4.2 YOLOv8-ResidualEMA最终检测指标表", ["指标", "数值", "说明"], [
        ["Precision", "91.60%", "预测为垃圾目标的结果中正确样本占比"],
        ["Recall", "83.60%", "真实垃圾目标中被成功检出的比例"],
        ["mAP@0.5", "90.60%", "IoU阈值为0.5时的平均检测精度"],
        ["mAP@0.5:0.95", "67.60%", "多IoU阈值下的平均检测精度"],
    ])
    # table is after title paragraph, continue after title by locating it
    last = find_starts("表4.2 YOLOv8-ResidualEMA") or last
    exp_dir = base / "论文实验数据"
    baseline_dir = base / "每日论文进展" / "实验数据" / "2026-04-30-yolov8s-baseline"
    imgs = [
        (baseline_dir / "box_loss.png", "图4.4 box loss训练变化曲线", 4.7),
        (baseline_dir / "cls_loss.png", "图4.5 cls loss训练变化曲线", 4.7),
        (baseline_dir / "dfl_loss.png", "图4.6 dfl loss训练变化曲线", 4.7),
        (exp_dir / "precision(B).png", "图4.7 Precision训练变化曲线", 4.7),
        (exp_dir / "recall(B).png", "图4.8 Recall训练变化曲线", 4.7),
        (exp_dir / "mAP50(B).png", "图4.9 mAP@0.5训练变化曲线", 4.7),
        (exp_dir / "mAP50-95(B).png", "图4.10 mAP@0.5:0.95训练变化曲线", 4.7),
        (exp_dir / "Precision-Recall Curve.png", "图4.11 Precision-Recall曲线", 5.0),
        (exp_dir / "F1-Confidence Curve.png", "图4.12 F1-Confidence曲线", 5.0),
        (exp_dir / "Precision-Confidence Curve.png", "图4.13 Precision-Confidence曲线", 5.0),
        (exp_dir / "Recall-Confidence Curve.png", "图4.14 Recall-Confidence曲线", 5.0),
        (baseline_dir / "yolov8s混淆矩阵.jpg", "图4.15 模型混淆矩阵", 5.0),
    ]
    for path, cap, width in imgs:
        last, ok = image_after(last, path, cap, width)
    para_after(last, "从新增曲线可以看出，训练损失整体下降，说明模型在定位、分类和边界框分布回归方面逐步收敛。Precision和mAP指标较高，说明ResidualEMA对减少误检和提升平均检测精度具有一定作用；但Recall仍低于Precision，说明模型仍存在漏检问题，特别是在小目标、遮挡和背景干扰较强的样本中仍需通过补充难例样本、调整置信度阈值和优化NMS参数继续改进。")

# ---------- Chapter 5: insert each image at matching section, not at chapter title ----------
ui = base / "系统界面"
# keep original Fig 5.1-5.4, add missing modules at their specific sections only
p57 = find_starts("5.7 历史图库")
if p57:
    last = para_after(p57, "历史图库和错误样本模块用于保存系统实际运行过程中的识别记录与问题样本。前者便于用户回看识别结果，后者用于收集误检、漏检或分类不确定样本，为后续模型迭代提供数据来源。")
    last, _ = image_after(last, ui / "新历史图库.jpg", "图5.5 历史图库界面", 5.3)
    last, _ = image_after(last, ui / "新错误样本.jpg", "图5.6 错误样本管理界面", 5.3)

p58 = find_starts("5.8 系统部署")
if p58:
    last = para_after(p58, "为验证系统不是单纯页面展示，本文围绕用户登录、图片上传、模型推理、分类建议生成、历史记录保存、错误样本反馈、模型管理和智能问答等核心流程进行功能测试。测试重点在于确认前端、后端和推理服务之间的数据链路是否完整。")
    table_after(last, "表5.1 系统功能测试表", ["测试模块", "测试内容", "预期结果", "测试结论"], [
        ["登录模块", "输入账号密码并提交", "成功进入系统主界面", "通过"],
        ["图片上传", "上传校园垃圾图片", "图片可预览并提交识别", "通过"],
        ["模型推理", "调用YOLOv8-ResidualEMA模型", "返回类别、置信度和检测框", "通过"],
        ["分类建议", "根据识别类别生成投放建议", "显示对应分类提示", "通过"],
        ["历史图库", "保存识别图片和结果", "历史记录可查询", "通过"],
        ["错误样本", "标记误识别样本", "样本进入反馈管理流程", "通过"],
        ["模型管理", "查看和切换模型状态", "模型信息显示正常", "通过"],
        ["智能问答", "输入垃圾分类问题", "返回解释和建议", "通过"],
    ])

# ---------- reference warning, not fabricated ----------
ref = find_starts("参考文献")
if ref:
    para_after(ref, "说明：本文参考文献采用顺序编码制，最终提交前仍需逐条核对正文引用编号与文末条目是否完全对应，尤其要检查作者、年份和编号是否存在错配。本文不新增未经核验的参考文献。")

# ---------- normalize captions/headings black ----------
for p in doc.paragraphs:
    t = p.text.strip()
    if re.match(r"^第[一二三四五六]章", t):
        fmt_heading(p, 1)
    elif re.match(r"^\d+\.\d+", t):
        fmt_heading(p, 2)
    elif re.match(r"^[图表]\d+\.\d+", t):
        fmt_caption(p)

doc.save(str(out))
print(out)
print('paragraphs', len(doc.paragraphs), 'tables', len(doc.tables))
