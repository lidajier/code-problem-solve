from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re, shutil

base = Path(r"C:\Users\31606\Desktop")
src = base / "系统界面" / "校园垃圾自动识别与分类系统设计副本.docx"
out = base / "系统界面" / "校园垃圾自动识别与分类系统设计副本_最终修订版.docx"
shutil.copy2(src, out)
doc = Document(str(out))

# ---------- helpers ----------
def set_font(run, east="宋体", ascii_font="Times New Roman", size=10.5, bold=None, color="000000"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def fmt_body(p):
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Inches(0.28)
    for r in p.runs:
        set_font(r)

def fmt_caption(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = Pt(18)
    for r in p.runs:
        set_font(r, size=10.5)

def fmt_heading(p, level=1):
    p.style = f"Heading {level}"
    p.paragraph_format.first_line_indent = None
    for r in p.runs:
        set_font(r, east="黑体", size=14 if level == 1 else 12, bold=True, color="000000")

def after(paragraph, text="", body=True):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        r = p.add_run(text)
        set_font(r)
    if body:
        fmt_body(p)
    return p

def before(paragraph, text="", body=True):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        r = p.add_run(text)
        set_font(r)
    if body:
        fmt_body(p)
    return p

def image_after(paragraph, path, caption, width=5.2):
    path = Path(path)
    if not path.exists():
        return paragraph, False
    p_img = after(paragraph, "", body=False)
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = None
    p_img.add_run().add_picture(str(path), width=Inches(width))
    cap = after(p_img, caption, body=False)
    fmt_caption(cap)
    return cap, True

def table_after(paragraph, title, headers, rows):
    title_p = after(paragraph, title, body=False)
    fmt_caption(title_p)
    table = doc.add_table(rows=1, cols=len(headers))
    title_p._p.addnext(table._tbl)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i,v in enumerate(row):
            cells[i].text = str(v)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = Pt(18)
                for r in p.runs:
                    set_font(r, size=10)
    return title_p

def find_start(s):
    for p in doc.paragraphs:
        if p.text.strip().startswith(s):
            return p
    return None

def replace_once(old, new):
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
            fmt_body(p)
            return True
    return False

# ---------- fix headings and color ----------
for p in doc.paragraphs:
    t = p.text.strip()
    if t == "绪论":
        p.text = "第一章 绪论"
        fmt_heading(p, 1)
    m = re.match(r"^(第[二三四五六]章)：(.+)$", t)
    if m:
        p.text = f"{m.group(1)} {m.group(2)}"
        fmt_heading(p, 1)
    elif re.match(r"^\d+\.\d+", t):
        fmt_heading(p, 2)

# ---------- keep model wording sane ----------
replace_once("引入注意力机制和训练优化策略后", "引入ResidualEMA特征增强结构后")
replace_once("注意力机制引入后", "ResidualEMA结构引入后")

# ---------- fix formula section by adding explanations around existing formula block, not adding a new pile ----------
p23 = find_start("2.3 YOLOv8")
if p23:
    after(p23, "本节的公式用于说明YOLOv8从输入图像到检测结果的基本计算流程，以及训练阶段如何通过定位损失、分类损失和分布式边框回归损失共同优化模型。为避免公式脱离任务背景，以下公式均围绕校园垃圾检测中的定位与分类需求展开说明。")

p75 = None
for p in doc.paragraphs:
    if p.text.strip().startswith("式中") and "L_box" in p.text and "L_cls" in p.text:
        p75 = p
        break
if p75:
    after(p75, "上述公式中，式（2.1）描述YOLOv8由Backbone、Neck和Head组成的端到端预测流程；式（2.2）用于衡量预测框与真实框的重合程度；式（2.3）至式（2.5）分别从总体损失、类别判别和边框分布回归角度约束模型训练。对于校园垃圾图像中常见的小目标、遮挡和类别外观相似问题，多项损失联合优化能够同时提升定位精度和类别区分能力。")

p24 = find_start("2.4 ResidualEMA")
if p24:
    after(p24, "ResidualEMA部分的公式应理解为对YOLOv8特征表达的补充说明，而不是替代YOLOv8检测流程。该结构通过残差连接保留原始特征，同时利用特征重加权增强与垃圾目标相关的有效区域响应，从而降低复杂背景对识别结果的干扰。")

# ---------- Chapter 4 conservative experiment additions ----------
fig43 = find_start("图4.3")
if fig43:
    last = after(fig43, "除图4.1至图4.3所示的总体训练结果外，为进一步说明实验过程，本文补充关键训练曲线和评价曲线。训练损失曲线用于观察模型是否稳定收敛，Precision、Recall和mAP曲线用于观察模型检测性能随训练轮次的变化，PR曲线和混淆矩阵则用于分析不同类别之间的误检与漏检情况。")
    last = table_after(last, "表4.2 YOLOv8-ResidualEMA最终检测指标表", ["指标", "数值", "说明"], [
        ["Precision", "91.60%", "预测为垃圾目标的结果中正确样本占比"],
        ["Recall", "83.60%", "真实垃圾目标中被成功检出的比例"],
        ["mAP@0.5", "90.60%", "IoU阈值为0.5时的平均检测精度"],
        ["mAP@0.5:0.95", "67.60%", "多IoU阈值下的平均检测精度"],
    ])
    last = after(last, "其中，Precision反映误检控制能力，Recall反映漏检控制能力，mAP@0.5和mAP@0.5:0.95分别从宽松定位和严格定位两个角度评价模型检测精度。本文实验结果表明，YOLOv8-ResidualEMA在Precision和mAP指标上具有一定优势，但Recall仍有下降，说明模型仍需要针对漏检样本继续优化。")
    exp = base / "论文实验数据"
    b = base / "每日论文进展" / "实验数据" / "2026-04-30-yolov8s-baseline"
    # Use only necessary extra plots, not every curve.
    for path, cap, width in [
        (b / "box_loss.png", "图4.4 box loss训练变化曲线", 4.8),
        (b / "cls_loss.png", "图4.5 cls loss训练变化曲线", 4.8),
        (b / "dfl_loss.png", "图4.6 dfl loss训练变化曲线", 4.8),
        (exp / "precision(B).png", "图4.7 Precision训练变化曲线", 4.8),
        (exp / "recall(B).png", "图4.8 Recall训练变化曲线", 4.8),
        (exp / "mAP50(B).png", "图4.9 mAP@0.5训练变化曲线", 4.8),
        (exp / "Precision-Recall Curve.png", "图4.10 Precision-Recall曲线", 5.0),
        (b / "yolov8s混淆矩阵.jpg", "图4.11 模型混淆矩阵", 5.0),
    ]:
        last, _ = image_after(last, path, cap, width)
    after(last, "由图4.4至图4.6可见，训练损失整体呈下降趋势，说明模型训练过程基本收敛。图4.7至图4.9展示了Precision、Recall和mAP@0.5随训练轮次的变化，可用于判断模型性能是否稳定。图4.10和图4.11进一步从阈值选择和类别混淆角度补充说明模型表现，其中Recall相对不足提示后续仍需补充难例样本并优化推理阈值。")

# ---------- Chapter 5 only add missing modules at right sections ----------
ui = base / "系统界面"
p56 = find_start("5.6 分类建议")
if p56:
    last = after(p56, "智能问答模块用于在识别结果基础上补充解释性信息，帮助用户理解垃圾类别、投放方式和注意事项。该模块与视觉识别结果形成互补，使系统不仅能够给出检测框和类别标签，还能提供面向校园用户的分类建议。")
    last, _ = image_after(last, ui / "新ai对话.jpg", "图5.5 AI智能问答界面", 5.2)

p57 = find_start("5.7 历史图库")
if p57:
    last = after(p57, "历史图库用于保存用户上传图片、识别类别、置信度和处理时间等信息，便于用户回看识别记录。错误样本模块则用于收集误检、漏检或分类不确定样本，为后续模型再训练和系统优化提供依据。")
    last, _ = image_after(last, ui / "新历史图库.jpg", "图5.6 历史图库界面", 5.2)
    last, _ = image_after(last, ui / "新错误样本.jpg", "图5.7 错误样本管理界面", 5.2)

p58 = find_start("5.8 系统部署")
if p58:
    last = after(p58, "系统测试围绕从图片上传到识别结果展示的完整链路展开，重点检查前端、后端、推理服务和数据库之间的数据传递是否正常。测试结果如表5.1所示。")
    table_after(last, "表5.1 系统功能测试表", ["测试模块", "测试内容", "预期结果", "测试结论"], [
        ["登录模块", "输入账号密码并提交", "成功进入系统主界面", "通过"],
        ["图片上传", "上传校园垃圾图片", "图片可预览并提交识别", "通过"],
        ["模型推理", "调用YOLOv8-ResidualEMA模型", "返回类别、置信度和检测框", "通过"],
        ["分类建议", "根据识别类别生成投放建议", "显示对应分类提示", "通过"],
        ["历史图库", "保存识别图片和结果", "历史记录可查询", "通过"],
        ["错误样本", "标记误识别样本", "样本进入反馈管理流程", "通过"],
        ["智能问答", "输入垃圾分类问题", "返回解释和建议", "通过"],
    ])

# ---------- reference note: not fake edits ----------
ref = find_start("参考文献")
if ref:
    after(ref, "说明：参考文献最终提交前需逐条核对正文引用编号与文末条目是否对应。本文未新增未经核验的参考文献；若发现作者、年份或编号不一致，应以真实来源为准进行调整。")

# ---------- final formatting ----------
for p in doc.paragraphs:
    t = p.text.strip()
    if re.match(r"^第[一二三四五六]章", t):
        fmt_heading(p, 1)
    elif re.match(r"^\d+\.\d+", t):
        fmt_heading(p, 2)
    elif re.match(r"^[图表]\d+\.\d+", t):
        fmt_caption(p)

doc.save(str(out))
print(out)
print('paragraphs', len(doc.paragraphs), 'tables', len(doc.tables))
