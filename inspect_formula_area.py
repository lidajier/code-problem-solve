from pathlib import Path
from docx import Document
p=Path(r"C:\Users\31606\Desktop\系统界面\校园垃圾自动识别与分类系统设计副本.docx")
d=Document(str(p))
for i in range(55, 95):
    if i < len(d.paragraphs):
        print(i, repr(d.paragraphs[i].text.strip()))
