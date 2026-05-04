from pathlib import Path
from docx import Document
from zipfile import ZipFile
p=Path(r"C:\Users\31606\Desktop\系统界面\校园垃圾自动识别与分类系统设计副本_最终修订版.docx")
orig=Path(r"C:\Users\31606\Desktop\系统界面\校园垃圾自动识别与分类系统设计副本.docx")
d=Document(str(p))
text='\n'.join(x.text for x in d.paragraphs)
print('final_exists', p.exists(), 'size', p.stat().st_size)
print('orig_size', orig.stat().st_size, 'orig_mtime', orig.stat().st_mtime)
with ZipFile(p) as z:
    print('zip_bad', z.testzip())
    print('media_count', len([n for n in z.namelist() if n.startswith('word/media/')]))
print('ResidualEMA', text.count('ResidualEMA'), 'FocalEMA', text.count('FocalEMA'), 'FocalModulation', text.count('FocalModulation'))
print('tables', len(d.tables), 'paragraphs', len(d.paragraphs))
print('--- positions ---')
for i,para in enumerate(d.paragraphs):
    t=para.text.strip()
    if t.startswith('第一章') or t.startswith('第四章') or t.startswith('第五章') or t.startswith('4.') or t.startswith('5.') or t.startswith('图4.') or t.startswith('图5.') or t.startswith('表4.') or t.startswith('表5.'):
        print(i, t)
